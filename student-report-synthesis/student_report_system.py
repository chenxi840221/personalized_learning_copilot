import os
import json
import uuid
import pandas as pd
from datetime import datetime
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.exceptions import ResourceNotFoundError
from openai import AzureOpenAI
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import docx
import tempfile
import subprocess
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

class StudentReportSystem:
    def __init__(self, form_recognizer_endpoint, form_recognizer_key, 
                 openai_endpoint, openai_key, openai_deployment):
        """Initialize the student report system with Azure service credentials."""
        # Set up logging
        self.logger = logging.getLogger(__name__)
        
        # Log initialization
        self.logger.info("Initializing StudentReportSystem")
        
        # Azure Document Intelligence (formerly Form Recognizer)
        self.document_client = None
        if form_recognizer_endpoint and form_recognizer_key:
            try:
                self.document_client = DocumentIntelligenceClient(
                    endpoint=form_recognizer_endpoint, 
                    credential=AzureKeyCredential(form_recognizer_key)
                )
                self.logger.info("Document Intelligence client initialized successfully")
            except Exception as e:
                self.logger.error(f"Failed to initialize Document Intelligence client: {e}")
        
        # Azure OpenAI
        self.openai_client = None
        if openai_endpoint and openai_key:
            try:
                self.openai_client = AzureOpenAI(
                    api_key=openai_key,  
                    api_version="2023-05-15",
                    azure_endpoint=openai_endpoint
                )
                self.logger.info("OpenAI client initialized successfully")
            except Exception as e:
                self.logger.error(f"Failed to initialize OpenAI client: {e}")
        
        self.openai_deployment = openai_deployment
        self.report_templates = {}
        self.report_structure = {}
        self.assessment_scales = {
            "achievement": {
                "A": "Outstanding",
                "B": "High",
                "C": "Expected",
                "D": "Basic",
                "E": "Limited"
            },
            "effort": {
                "High": "Consistently applies themselves",
                "Satisfactory": "Generally applies themselves",
                "Low": "Inconsistently applies themselves"
            }
        }
        
        # Check if LibreOffice is installed for Word to PDF conversion
        self.libreoffice_path = self._find_libreoffice()
        if self.libreoffice_path:
            self.logger.info(f"LibreOffice found at: {self.libreoffice_path}")
        else:
            self.logger.warning("LibreOffice not found. Word to PDF conversion might be limited.")
    
    def _find_libreoffice(self):
        """Find the LibreOffice executable on the system."""
        possible_paths = [
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/usr/lib/libreoffice/program/soffice",
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        # Try finding it in PATH
        try:
            result = subprocess.run(['which', 'soffice'], capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass
            
        # Try with where command on Windows
        try:
            result = subprocess.run(['where', 'soffice'], capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass
        
        return None
    
    def convert_word_to_pdf(self, word_path):
        """Convert a Word document to PDF for processing."""
        self.logger.info(f"Converting Word document to PDF: {word_path}")
        
        # Create a temporary file for the PDF output
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
            pdf_path = tmp_pdf.name
        
        try:
            if self.libreoffice_path:
                # Convert using LibreOffice
                output_dir = os.path.dirname(pdf_path)
                cmd = [
                    self.libreoffice_path,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    word_path
                ]
                
                self.logger.debug(f"Running conversion command: {' '.join(cmd)}")
                process = subprocess.run(cmd, capture_output=True, text=True)
                
                if process.returncode != 0:
                    self.logger.error(f"LibreOffice conversion failed: {process.stderr}")
                    raise Exception(f"Failed to convert Word to PDF: {process.stderr}")
                
                # LibreOffice creates the PDF with the same basename in the output directory
                base_name = os.path.basename(word_path)
                base_name_without_ext = os.path.splitext(base_name)[0]
                new_pdf_path = os.path.join(output_dir, f"{base_name_without_ext}.pdf")
                
                # Rename to our expected path if needed
                if new_pdf_path != pdf_path and os.path.exists(new_pdf_path):
                    os.replace(new_pdf_path, pdf_path)
            else:
                # Fallback to python-docx and ReportLab if LibreOffice is not available
                self.logger.info("Using python-docx fallback for conversion")
                self._convert_with_python_docx(word_path, pdf_path)
            
            self.logger.info(f"Word document converted to PDF at: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            self.logger.error(f"Error converting Word to PDF: {e}")
            # If conversion fails, return None
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
            return None
    
    def _convert_with_python_docx(self, word_path, pdf_path):
        """Fallback method to convert Word to PDF using python-docx and ReportLab."""
        try:
            # Load the Word document
            doc = docx.Document(word_path)
            
            # Create a PDF document
            pdf = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            heading1_style = styles['Heading1']
            heading2_style = styles['Heading2']
            
            # Process the document
            elements = []
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                # Check paragraph style and apply appropriate PDF style
                if paragraph.style.name.startswith('Heading1') or paragraph.style.name == 'Title':
                    elements.append(Paragraph(paragraph.text, heading1_style))
                elif paragraph.style.name.startswith('Heading'):
                    elements.append(Paragraph(paragraph.text, heading2_style))
                else:
                    elements.append(Paragraph(paragraph.text, normal_style))
                
                elements.append(Spacer(1, 12))
            
            # Process tables (basic support)
            for table in doc.tables:
                data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Get text from cell paragraphs
                        text = '\n'.join(p.text for p in cell.paragraphs)
                        row_data.append(text)
                    data.append(row_data)
                
                if data:
                    pdf_table = Table(data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ]))
                    elements.append(pdf_table)
                    elements.append(Spacer(1, 12))
            
            # Build the PDF
            pdf.build(elements)
            
        except Exception as e:
            self.logger.error(f"Error in python-docx conversion: {e}")
            raise
    
    def extract_template_structure(self, template_path, template_name):
        """Extract structure from a report template using Azure Document Intelligence."""
        self.logger.info(f"Processing template: {template_name} from {template_path}")
        
        if not self.document_client:
            self.logger.error("Document client not initialized. Cannot extract template structure.")
            return {
                "error": "Document Intelligence client not initialized",
                "status": "failed"
            }
        
        # Check if the file is a Word document and convert if needed
        if template_path.lower().endswith(('.docx', '.doc')):
            self.logger.info(f"Template is a Word document: {template_path}")
            converted_pdf_path = self.convert_word_to_pdf(template_path)
            if not converted_pdf_path:
                return {
                    "error": "Failed to convert Word document to PDF",
                    "status": "failed"
                }
            template_path = converted_pdf_path
            self.logger.info(f"Using converted PDF: {template_path}")
        
        try:
            with open(template_path, "rb") as template_file:
                poller = self.document_client.begin_analyze_document(
                    "prebuilt-layout", template_file
                )
                result = poller.result()
                
            # Process and store the template structure
            template_data = {
                "pages": [],
                "tables": [],
                "headings": [],
                "paragraphs": []
            }
            
            for page in result.pages:
                page_data = {
                    "page_number": page.page_number,
                    "width": page.width,
                    "height": page.height,
                    "text_sections": []
                }
                
                for line in page.lines:
                    page_data["text_sections"].append({
                        "text": line.content,
                        "bounding_box": line.polygon,
                        "position": (line.polygon[0], line.polygon[1])
                    })
                    
                    # Identify headings based on formatting
                    if len(line.content) < 100 and line.content.strip().endswith(':') or line.content.isupper():
                        template_data["headings"].append(line.content)
                
                template_data["pages"].append(page_data)
                
            # Extract tables from the document
            for table in result.tables:
                table_data = {
                    "row_count": table.row_count,
                    "column_count": table.column_count,
                    "cells": []
                }
                
                for cell in table.cells:
                    table_data["cells"].append({
                        "row_index": cell.row_index,
                        "column_index": cell.column_index,
                        "content": cell.content,
                        "is_header": cell.kind == "columnHeader" or cell.kind == "rowHeader"
                    })
                    
                template_data["tables"].append(table_data)
            
            # Add template metadata
            template_data["name"] = template_name
            template_data["extraction_date"] = datetime.now().isoformat()
            template_data["status"] = "completed"
            
            # Store the extracted template structure
            self.report_templates[template_name] = template_data
            
            self.logger.info(f"Successfully extracted template structure for: {template_name}")
            
            # Clean up temporary PDF file if we converted from Word
            if template_path.endswith('.pdf') and template_path != template_path:
                try:
                    os.unlink(template_path)
                except Exception as e:
                    self.logger.warning(f"Could not delete temporary PDF: {e}")
            
            return template_data
            
        except Exception as e:
            self.logger.error(f"Error extracting template structure: {e}")
            return {
                "error": str(e),
                "status": "failed"
            }
    
    def analyze_report_structure(self):
        """Use Azure OpenAI to analyze the report templates and create a unified structure."""
        
        if not self.openai_client:
            self.logger.error("OpenAI client not initialized. Cannot analyze report structure.")
            return None
        
        if not self.report_templates:
            self.logger.error("No templates available for analysis.")
            return None
        
        self.logger.info("Analyzing report structure from templates...")
        
        # Combine all template data for analysis
        template_data_text = json.dumps(self.report_templates, indent=2)
        
        # Request OpenAI to analyze the structure and provide a unified format
        system_prompt = """
        You are an education document specialist tasked with analyzing student report templates.
        Extract the common structural elements, content sections, and formatting guidelines.
        Identify assessment scales, grading systems, and educational terminology.
        Create a unified report structure that follows Australian educational guidelines.
        """
        
        user_prompt = f"""
        Analyze these student report templates and extract:
        1. Common sections and their order
        2. Assessment scales and descriptions
        3. Required content for each section
        4. Formatting guidelines
        
        Template data:
        {template_data_text}
        
        Return a JSON structure defining a unified report template compatible with the input templates.
        """
        
        try:
            response = self.openai_client.chat.completions.create(
                model=self.openai_deployment,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.2,
                max_tokens=2000
            )
            
            # Extract and parse the unified structure
            result_text = response.choices[0].message.content
            self.logger.debug(f"OpenAI response for structure analysis: {result_text}")
            
            try:
                # Extract JSON from response if it's wrapped in markdown code blocks
                if "```json" in result_text:
                    json_text = result_text.split("```json")[1].split("```")[0].strip()
                    self.report_structure = json.loads(json_text)
                else:
                    self.report_structure = json.loads(result_text)
                    
                # Extract assessment scales for later use
                if "assessment_scales" in self.report_structure:
                    self.assessment_scales = self.report_structure["assessment_scales"]
                
                self.logger.info("Successfully analyzed report structure")
                return self.report_structure
                
            except json.JSONDecodeError as e:
                self.logger.error(f"Error parsing OpenAI response: {e}")
                self.logger.debug(f"Response text: {result_text}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error analyzing report structure: {e}")
            return None
    
    def generate_student_data(self, count=1):
        """Generate synthetic student data for report generation."""
        if not self.openai_client:
            self.logger.error("OpenAI client not initialized. Cannot generate student data.")
            return []
        
        self.logger.info(f"Generating synthetic data for {count} students")
        
        # Request OpenAI to generate synthetic student data
        system_prompt = """
        You are an education data specialist who creates realistic student data for testing reporting systems.
        Create realistic, varied student profiles with appropriate academic performance, social development, 
        and learning characteristics for Australian primary school students.
        """
        
        user_prompt = f"""
        Create {count} synthetic student profiles for primary school students (ages 5-12).
        Include:
        - Full name, age, grade level
        - Academic performance in key subjects (English, Mathematics, Science, etc.)
        - Social development and behavior traits
        - Learning strengths and areas for improvement
        - Attendance data
        
        Use the assessment scales from our system:
        {json.dumps(self.assessment_scales, indent=2)}
        
        Return data in JSON format that could be used to populate student reports.
        """
        
        try:
            response = self.openai_client.chat.completions.create(
                model=self.openai_deployment,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            
            # Extract and parse the student data
            result_text = response.choices[0].message.content
            self.logger.debug(f"OpenAI response for student data generation: {result_text}")
            
            try:
                # Extract JSON from response if it's wrapped in markdown code blocks
                if "```json" in result_text:
                    json_text = result_text.split("```json")[1].split("```")[0].strip()
                    students = json.loads(json_text)
                else:
                    students = json.loads(result_text)
                    
                self.logger.info(f"Successfully generated data for {len(students)} students")
                return students
                
            except json.JSONDecodeError as e:
                self.logger.error(f"Error parsing OpenAI response for student data: {e}")
                self.logger.debug(f"Response text: {result_text}")
                return []
                
        except Exception as e:
            self.logger.error(f"Error generating student data: {e}")
            return []
    
    def generate_student_report(self, student_data):
        """Generate a student report based on student data and the unified report structure."""
        if not self.openai_client:
            self.logger.error("OpenAI client not initialized. Cannot generate student report.")
            return None
        
        self.logger.info(f"Generating report for student: {student_data.get('full_name', 'Unknown')}")
        
        # If report structure hasn't been analyzed yet, use a default structure
        if not self.report_structure:
            self.report_structure = {
                "sections": [
                    "student_info",
                    "general_comment",
                    "academic_subjects",
                    "social_development",
                    "attendance"
                ],
                "academic_subjects": [
                    "English",
                    "Mathematics",
                    "Science",
                    "Humanities",
                    "Arts",
                    "Physical Education"
                ]
            }
        
        # Request OpenAI to generate the report content
        system_prompt = """
        You are an experienced teacher writing student reports according to Australian educational guidelines.
        Write professionally using clear, concise language focusing on the student's achievements and areas for growth.
        Avoid jargon, be specific rather than general, and provide constructive feedback.
        Your reports should be informative, positive, and helpful to both students and parents.
        """
        
        user_prompt = f"""
        Generate a complete student report for the following student:
        {json.dumps(student_data, indent=2)}
        
        Follow this report structure:
        {json.dumps(self.report_structure, indent=2)}
        
        For each section in the report structure, provide appropriate, professionally-written content.
        Return the complete report in JSON format matching our report structure.
        For each assessment area, include both the grade/rating and a written comment.
        """
        
        try:
            response = self.openai_client.chat.completions.create(
                model=self.openai_deployment,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.5,
                max_tokens=3000
            )
            
            # Extract and parse the report content
            result_text = response.choices[0].message.content
            self.logger.debug(f"OpenAI response for report generation: {result_text}")
            
            try:
                # Extract JSON from response if it's wrapped in markdown code blocks
                if "```json" in result_text:
                    json_text = result_text.split("```json")[1].split("```")[0].strip()
                    report_content = json.loads(json_text)
                else:
                    report_content = json.loads(result_text)
                    
                self.logger.info(f"Successfully generated report content")
                return report_content
                
            except json.JSONDecodeError as e:
                self.logger.error(f"Error parsing OpenAI response for report: {e}")
                self.logger.debug(f"Response text: {result_text}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error generating report: {e}")
            return None
    
    def create_pdf_report(self, report_content, output_path):
        """Create a PDF report based on the generated report content."""
        self.logger.info(f"Creating PDF report at: {output_path}")
        
        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Create styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                fontSize=16,
                alignment=1,  # Center
                spaceAfter=12
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=6
            )
            
            subheading_style = ParagraphStyle(
                'Subheading',
                parent=styles['Heading3'],
                fontSize=12,
                spaceAfter=6
            )
            
            normal_style = styles['Normal']
            
            # Build the report content
            story = []
            
            # School header
            if "school_info" in report_content:
                school = report_content["school_info"]
                school_name = school.get("name", "School Name")
                story.append(Paragraph(school_name, title_style))
                story.append(Spacer(1, 12))
                
                # School contact info
                if "contact_info" in school:
                    contact = school["contact_info"]
                    contact_text = f"Email: {contact.get('email', '')}, Phone: {contact.get('phone', '')}"
                    story.append(Paragraph(contact_text, normal_style))
                    story.append(Spacer(1, 12))
            else:
                # Default school info if not provided
                story.append(Paragraph("Australian Primary School", title_style))
                story.append(Spacer(1, 12))
            
            # Student information
            if "student_info" in report_content:
                student = report_content["student_info"]
                student_name = student.get("name", student.get("full_name", "Student Name"))
                student_grade = student.get("grade", student.get("grade_level", ""))
                story.append(Paragraph(f"Student Report: {student_name}", heading_style))
                story.append(Paragraph(f"Grade: {student_grade}", normal_style))
                
                # Term/Year
                term = report_content.get("term", "1")
                year = report_content.get("year", datetime.now().year)
                story.append(Paragraph(f"Term: {term}, Year: {year}", normal_style))
                story.append(Spacer(1, 12))
            
            # General comment
            if "general_comment" in report_content:
                story.append(Paragraph("General Comment", heading_style))
                story.append(Paragraph(report_content["general_comment"], normal_style))
                story.append(Spacer(1, 18))
            
            # Social development
            if "social_development" in report_content:
                story.append(Paragraph("Social Development and Commitment to Learning", heading_style))
                
                # Create a table for social development ratings
                if isinstance(report_content["social_development"], list):
                    data = [["Area", "Rating"]]
                    for item in report_content["social_development"]:
                        data.append([item.get("area", ""), item.get("rating", "")])
                        
                    t = Table(data, colWidths=[300, 100])
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(t)
                else:
                    story.append(Paragraph(report_content["social_development"], normal_style))
                
                story.append(Spacer(1, 18))
            
            # Academic subjects
            if "academic_subjects" in report_content:
                story.append(Paragraph("Academic Progress", heading_style))
                
                for subject in report_content["academic_subjects"]:
                    subject_name = subject.get("name", "Subject")
                    story.append(Paragraph(subject_name, subheading_style))
                    
                    # Create a table for achievement and effort
                    data = [["Achievement", "Effort"]]
                    data.append([subject.get("achievement", ""), subject.get("effort", "")])
                    
                    t = Table(data, colWidths=[200, 200])
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 6))
                    
                    # Teacher's comments
                    if "comment" in subject:
                        story.append(Paragraph("Teacher's Comments:", normal_style))
                        story.append(Paragraph(subject["comment"], normal_style))
                    
                    story.append(Spacer(1, 12))
            
            # Attendance
            if "attendance" in report_content:
                story.append(Paragraph("Attendance", heading_style))
                
                attendance = report_content["attendance"]
                data = [["Days Absent", "Partial Days Absent"]]
                data.append([
                    str(attendance.get("days_absent", 0)), 
                    str(attendance.get("partial_days_absent", 0))
                ])
                
                t = Table(data, colWidths=[200, 200])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
                story.append(Spacer(1, 18))
            
            # Extra activities
            if "extra_activities" in report_content and report_content["extra_activities"]:
                story.append(Paragraph("Extra Activities", heading_style))
                
                activities = report_content["extra_activities"]
                if isinstance(activities, list):
                    for activity in activities:
                        story.append(Paragraph(f"• {activity}", normal_style))
                else:
                    story.append(Paragraph(activities, normal_style))
                
                story.append(Spacer(1, 18))
            
            # Teacher signature
            if "teacher_info" in report_content:
                teacher = report_content["teacher_info"]
                story.append(Paragraph(f"Teacher: {teacher.get('name', '')}", normal_style))
                story.append(Paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}", normal_style))
            else:
                # Default teacher info if not provided
                story.append(Paragraph(f"Teacher: Class Teacher", normal_style))
                story.append(Paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}", normal_style))
            
            # Make sure the output directory exists
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            
            # Build the PDF
            doc.build(story)
            
            # Verify the file was created
            if not os.path.exists(output_path):
                self.logger.error(f"PDF creation failed, file does not exist: {output_path}")
                return None
                
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                self.logger.error(f"PDF creation failed, file is empty: {output_path}")
                return None
                
            self.logger.info(f"Successfully created PDF report at: {output_path} ({file_size} bytes)")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error creating PDF report: {e}")
            return None